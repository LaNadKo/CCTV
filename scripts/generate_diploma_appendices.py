from __future__ import annotations

import math
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Mm, Pt


PROJECT_ROOT = Path(r"C:\Users\dsok8\Desktop\CCTV Комплекс")
DIPLOMA_DIR = Path(r"C:\Users\dsok8\Desktop\Диплом")

MAX_CODE_COLUMNS = 92
MAX_VISUAL_LINES_PER_PAGE = 58
HEADER_COST = 4
CAPTION_COST = 3
BLOCK_OVERHEAD = 2
MIN_FREE_LINES_FOR_NEW_BLOCK = 8


@dataclass(frozen=True)
class Appendix:
    number: int
    title: str
    output_name: str
    files: list[Path]
    caption_root: Path


def ensure_style(doc: Document, style_name: str, fallback: str) -> str:
    try:
        doc.styles[style_name]
        return style_name
    except KeyError:
        return fallback


def set_paragraph_pagination(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


def apply_times(paragraph, *, bold: bool, size: int) -> None:
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        run.font.size = Pt(size)
        run.bold = bold


def add_page_header(doc: Document, text: str, *, first_page: bool, heading_style: str, normal_style: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = heading_style if first_page else normal_style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.first_line_indent = Cm(1.25)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run(text)
    apply_times(paragraph, bold=first_page, size=14)
    set_paragraph_pagination(paragraph)


def sanitize_xml_text(text: str) -> str:
    return "".join(ch for ch in text if ch in ("\t", "\n", "\r") or ord(ch) >= 32)


def add_code_table(doc: Document, code_text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.autofit = True
    cell = table.cell(0, 0)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(sanitize_xml_text(code_text))
    run.font.name = "Courier New"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
    run.font.size = Pt(10)


def add_caption(doc: Document, figure_number: int, relative_path: str, normal_style: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.style = normal_style
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run(f"Рисунок {figure_number} - {relative_path}")
    apply_times(paragraph, bold=False, size=14)


def configure_document(doc: Document) -> tuple[str, str]:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(30)
    section.right_margin = Mm(15)
    return ensure_style(doc, "Heading 1", "Normal"), ensure_style(doc, "Normal", "Normal")


def remove_initial_empty_paragraph(doc: Document) -> None:
    if len(doc.paragraphs) != 1:
        return
    paragraph = doc.paragraphs[0]
    if paragraph.text.strip():
        return
    paragraph._element.getparent().remove(paragraph._element)


def visual_line_cost(line: str) -> int:
    expanded = line.expandtabs(4)
    width = max(1, len(expanded))
    return max(1, math.ceil(width / MAX_CODE_COLUMNS))


def chunk_lines(lines: list[str], max_visual_lines: int) -> tuple[list[str], int]:
    if not lines:
        return [], 0

    taken: list[str] = []
    used = 0
    for line in lines:
        cost = visual_line_cost(line)
        if taken and used + cost > max_visual_lines:
            break
        if not taken and cost > max_visual_lines:
            taken.append(line)
            used += cost
            break
        taken.append(line)
        used += cost
    return taken, used


def read_text_lines(path: Path) -> list[str]:
    for encoding in ("utf-8-sig", "utf-8", "cp1251"):
        try:
            return path.read_text(encoding=encoding).splitlines()
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore").splitlines()


def relative_caption(path: Path, root: Path) -> str:
    try:
        return str(path.relative_to(root)).replace("/", "\\")
    except ValueError:
        return str(path).replace("/", "\\")


def unique_existing(paths: list[Path]) -> list[Path]:
    result: list[Path] = []
    seen: set[str] = set()
    for path in paths:
        if not path.exists() or not path.is_file():
            continue
        try:
            if path.stat().st_size == 0:
                continue
        except OSError:
            continue
        key = str(path.resolve()).lower()
        if key in seen:
            continue
        seen.add(key)
        result.append(path)
    return result


def collect_by_suffix(base: Path, suffixes: set[str]) -> list[Path]:
    if not base.exists():
        return []
    paths: list[Path] = []
    for path in sorted(base.rglob("*")):
        if path.is_file() and path.suffix.lower() in suffixes:
            paths.append(path)
    return paths


def collect_program_files() -> list[Path]:
    root_files = [
        ".dockerignore",
        ".env.example",
        ".gitignore",
        "alembic.ini",
        "docker-compose.yml",
        "docker-entrypoint.sh",
        "Dockerfile",
        "init-ssl.sh",
        "install.sh",
        "requirements.txt",
        "start-all-docker.bat",
        "start-all.bat",
        "start-backend.bat",
        "start-db.bat",
        "start-local.bat",
        "stop-all.bat",
        "stop-local.bat",
    ]
    files = [PROJECT_ROOT / item for item in root_files]

    files += collect_by_suffix(PROJECT_ROOT / "app", {".py"})
    files += collect_by_suffix(PROJECT_ROOT / "cctv_ai", {".py"})
    files += collect_by_suffix(PROJECT_ROOT / "migrations", {".py"})
    files += collect_by_suffix(PROJECT_ROOT / "processor", {".py", ".txt", ".yml", ".yaml", ".sh", ".bat", ".iss", ".service", ".example", ".spec"})

    for app_dir in ("native/cctv_console", "native/cctv_processor_gui"):
        base = PROJECT_ROOT / app_dir
        files += [base / "pubspec.yaml", base / "analysis_options.yaml"]
        files += collect_by_suffix(base / "lib", {".dart"})

    files += collect_by_suffix(PROJECT_ROOT / "nginx", {".conf", ".template", ".yml", ".yaml"})

    excluded_names = {"package-lock.json", "pubspec.lock", "processor.log"}
    excluded_parts = {
        ".git",
        "__pycache__",
        "node_modules",
        ".dart_tool",
        "build",
        "dist",
        "release",
        "recordings",
        "recordings_cache",
        "snapshots",
        "data",
        "assets",
        "test",
    }

    filtered: list[Path] = []
    for path in files:
        rel_parts = set(path.relative_to(PROJECT_ROOT).parts) if path.is_relative_to(PROJECT_ROOT) else set(path.parts)
        if path.name in excluded_names:
            continue
        if excluded_parts & rel_parts:
            continue
        filtered.append(path)

    return unique_existing(filtered)


def collect_testing_files() -> list[Path]:
    files = [
        PROJECT_ROOT / "scripts" / "smoke_api.py",
    ]
    files += collect_by_suffix(PROJECT_ROOT / "native" / "cctv_console" / "test", {".dart"})
    files += collect_by_suffix(PROJECT_ROOT / "native" / "cctv_processor_gui" / "test", {".dart"})
    return unique_existing(files)


def build_document(appendix: Appendix) -> Document:
    doc = Document()
    heading_style, normal_style = configure_document(doc)
    remove_initial_empty_paragraph(doc)

    figure_number = 1
    started = False
    page_remaining = 0

    for path in appendix.files:
        lines = read_text_lines(path)
        if not lines:
            continue

        remaining = list(lines)
        while remaining:
            if not started or page_remaining == 0:
                if started:
                    doc.add_page_break()
                header = (
                    f"ПРИЛОЖЕНИЕ {appendix.number} - {appendix.title}"
                    if not started
                    else f"ПРОДОЛЖЕНИЕ ПРИЛОЖЕНИЯ {appendix.number}"
                )
                add_page_header(
                    doc,
                    header,
                    first_page=not started,
                    heading_style=heading_style,
                    normal_style=normal_style,
                )
                started = True
                page_remaining = MAX_VISUAL_LINES_PER_PAGE - HEADER_COST

            available_for_code = max(1, page_remaining - CAPTION_COST - BLOCK_OVERHEAD)
            chunk, used_lines = chunk_lines(remaining, available_for_code)
            add_code_table(doc, "\n".join(chunk))
            add_caption(doc, figure_number, relative_caption(path, appendix.caption_root), normal_style)
            figure_number += 1
            remaining = remaining[len(chunk):]
            page_remaining -= used_lines + CAPTION_COST + BLOCK_OVERHEAD

            if page_remaining < MIN_FREE_LINES_FOR_NEW_BLOCK:
                page_remaining = 0

    return doc


def main() -> None:
    appendices = [
        Appendix(
            number=9,
            title="SQL-СЦЕНАРИЙ",
            output_name="ПРИЛОЖЕНИЕ 9 - SQL-СЦЕНАРИЙ.docx",
            files=unique_existing([DIPLOMA_DIR / "cctv_actual_schema_english.sql"]),
            caption_root=DIPLOMA_DIR,
        ),
        Appendix(
            number=10,
            title="ЛИСТИНГ ПРОГРАММЫ",
            output_name="ПРИЛОЖЕНИЕ 10 - ЛИСТИНГ ПРОГРАММЫ.docx",
            files=collect_program_files(),
            caption_root=PROJECT_ROOT,
        ),
        Appendix(
            number=11,
            title="ЛИСТИНГ ТЕСТИРОВАНИЯ",
            output_name="ПРИЛОЖЕНИЕ 11 - ЛИСТИНГ ТЕСТИРОВАНИЯ.docx",
            files=collect_testing_files(),
            caption_root=PROJECT_ROOT,
        ),
    ]

    DIPLOMA_DIR.mkdir(parents=True, exist_ok=True)
    for appendix in appendices:
        doc = build_document(appendix)
        output_path = DIPLOMA_DIR / appendix.output_name
        doc.save(output_path)
        print(f"{appendix.output_name}: {len(appendix.files)} files -> {output_path}")


if __name__ == "__main__":
    main()
