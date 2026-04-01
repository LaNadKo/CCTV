from __future__ import annotations

import json
from collections import Counter, defaultdict
from datetime import datetime, timedelta, timezone
from io import BytesIO
from pathlib import Path
from typing import Optional, Tuple

from fastapi import APIRouter, Depends, HTTPException, Query, status
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app import models
from app.db import get_session
from app.dependencies import get_current_user
from app.permissions import is_at_least_user
from app.schemas.reports import (
    AppearanceItem,
    AppearanceReport,
    ArchiveCameraStat,
    ArchiveReport,
    CameraReportItem,
    EventReviewReport,
    GroupReportItem,
    ProcessorReportItem,
    RecentUserAction,
    ReportsDashboard,
    ReportStorageStat,
    ReportValueLabel,
    ReviewerStat,
    SecurityFailureItem,
    SecurityReport,
    UserActionActorStat,
    UserActionsReport,
)

router = APIRouter(prefix="/reports", tags=["reports"])


def _parse_iso_datetime(value: Optional[str]) -> Optional[datetime]:
    if not value:
        return None
    try:
        normalized = value.replace("Z", "+00:00")
        parsed = datetime.fromisoformat(normalized)
        if parsed.tzinfo is not None:
            parsed = parsed.astimezone(timezone.utc).replace(tzinfo=None)
        return parsed
    except ValueError:
        return None


def _format_iso(value: Optional[datetime]) -> Optional[str]:
    return value.isoformat() if value else None


def _format_ts(value: Optional[str]) -> str:
    if not value:
        return ""
    try:
        normalized = value.replace("Z", "+00:00")
        return datetime.fromisoformat(normalized).strftime("%d.%m.%Y %H:%M:%S")
    except Exception:
        return value


def _person_label(person: Optional[models.Person]) -> Optional[str]:
    if person is None:
        return None
    parts = [person.last_name, person.first_name, person.middle_name]
    return " ".join([part for part in parts if part]) or f"ID {person.person_id}"


def _user_label(user: Optional[models.User], user_id: Optional[int] = None) -> str:
    if user is None:
        return f"ID {user_id}" if user_id else "Система"
    parts = [user.last_name, user.first_name, user.middle_name]
    return " ".join([part for part in parts if part]) or user.login or f"ID {user.user_id}"


def _format_period(date_from: Optional[str], date_to: Optional[str]) -> str:
    if date_from and date_to:
        return f"{date_from} - {date_to}"
    if date_from:
        return f"с {date_from}"
    if date_to:
        return f"по {date_to}"
    return "за весь период"


def _resolve_font_path() -> Optional[str]:
    bundled = Path(__file__).resolve().parents[1] / "assets" / "fonts" / "Roboto-Regular.ttf"
    if bundled.exists():
        return str(bundled)
    for candidate in (
        Path("C:/Windows/Fonts/arial.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    ):
        if candidate.exists():
            return str(candidate)
    return None


def _safe_int(value) -> int:
    if value is None:
        return 0
    try:
        return int(value)
    except Exception:
        return 0


def _format_bytes(value: Optional[int]) -> str:
    amount = float(value or 0)
    if amount <= 0:
        return "0 Б"
    units = ["Б", "КБ", "МБ", "ГБ", "ТБ"]
    index = 0
    while amount >= 1024 and index < len(units) - 1:
        amount /= 1024
        index += 1
    precision = 0 if amount >= 10 or index == 0 else 1
    return f"{amount:.{precision}f} {units[index]}"


def _format_duration(seconds: Optional[float]) -> str:
    if not seconds or seconds <= 0:
        return "-"
    total = round(seconds)
    hours = total // 3600
    minutes = (total % 3600) // 60
    secs = total % 60
    if hours:
        return f"{hours} ч {minutes} мин"
    if minutes:
        return f"{minutes} мин {secs} с"
    return f"{secs} с"


def _is_processor_online(processor: models.Processor) -> bool:
    if processor.status != "online" or not processor.last_heartbeat:
        return False
    return (datetime.utcnow() - processor.last_heartbeat) <= timedelta(seconds=120)


def _load_metrics(processor: models.Processor) -> dict:
    if not processor.last_metrics:
        return {}
    try:
        return json.loads(processor.last_metrics)
    except Exception:
        return {}


async def _load_appearance_items(
    date_from: Optional[str],
    date_to: Optional[str],
    person_id: Optional[int],
    session: AsyncSession,
) -> Tuple[list[AppearanceItem], Optional[str], Optional[str], Optional[str]]:
    stmt = (
        select(models.Event, models.Person, models.Camera, models.Group)
        .join(models.EventType, models.Event.event_type_id == models.EventType.event_type_id)
        .outerjoin(models.Person, models.Event.person_id == models.Person.person_id)
        .outerjoin(models.Camera, models.Event.camera_id == models.Camera.camera_id)
        .outerjoin(models.Group, models.Camera.group_id == models.Group.group_id)
        .where(models.EventType.name == "face_recognized")
        .order_by(models.Event.event_ts.asc())
    )

    if person_id is not None:
        stmt = stmt.where(models.Event.person_id == person_id)

    date_from_dt = _parse_iso_datetime(date_from)
    date_to_dt = _parse_iso_datetime(date_to)
    if date_from_dt:
        stmt = stmt.where(models.Event.event_ts >= date_from_dt)
    if date_to_dt:
        stmt = stmt.where(models.Event.event_ts <= date_to_dt)

    rows = (await session.execute(stmt)).all()
    items: list[AppearanceItem] = []
    for event, person, camera, group in rows:
        items.append(
            AppearanceItem(
                event_id=event.event_id,
                event_ts=event.event_ts.isoformat(),
                camera_id=event.camera_id,
                camera_name=camera.name if camera else None,
                camera_location=camera.location if camera else None,
                group_name=group.name if group else None,
                person_id=event.person_id,
                person_label=_person_label(person),
                confidence=float(event.confidence) if event.confidence is not None else None,
            )
        )

    person_label = None
    if person_id is not None:
        person = await session.get(models.Person, person_id)
        person_label = _person_label(person) if person else f"ID {person_id}"

    return items, person_label, date_from, date_to


def _appearance_row(index: int, item: AppearanceItem) -> list[str]:
    return [
        str(index),
        _format_ts(item.event_ts),
        item.camera_name or f"Камера {item.camera_id}",
        item.camera_location or "-",
        item.group_name or "-",
        item.person_label or (f"ID {item.person_id}" if item.person_id else "-"),
        f"{item.confidence:.2f}" if item.confidence is not None else "-",
    ]


def _measure_columns(headers: list[str], rows: list[list[str]], min_width: int = 8, max_width: int = 42) -> list[int]:
    widths: list[int] = []
    for index, header in enumerate(headers):
        max_len = len(str(header))
        for row in rows:
            if index < len(row):
                max_len = max(max_len, len(str(row[index])))
        widths.append(min(max(int(max_len * 1.08) + 2, min_width), max_width))
    return widths


def _fit_pdf_widths(headers: list[str], rows: list[list[str]], available_width: float) -> list[float]:
    raw = _measure_columns(headers, rows, min_width=8, max_width=36)
    total = max(sum(raw), 1)
    return [available_width * (width / total) for width in raw]


def _render_export_table(
    *,
    title: str,
    summary_lines: list[str],
    headers: list[str],
    rows: list[list[str]],
    fmt: str,
    filename_prefix: str,
):
    filename = f"{filename_prefix}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{fmt}"
    column_widths = _measure_columns(headers, rows)

    if fmt == "xlsx":
        from openpyxl import Workbook
        from openpyxl.utils import get_column_letter
        from openpyxl.styles import Alignment, Font, PatternFill

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Отчет"
        sheet["A1"] = title
        sheet.merge_cells(start_row=1, start_column=1, end_row=1, end_column=len(headers))
        sheet["A1"].font = Font(size=14, bold=True)
        sheet["A1"].alignment = Alignment(horizontal="center")

        row_index = 2
        for line in summary_lines:
            sheet.cell(row=row_index, column=1, value=line)
            sheet.merge_cells(start_row=row_index, start_column=1, end_row=row_index, end_column=len(headers))
            row_index += 1

        for index, header in enumerate(headers, start=1):
            cell = sheet.cell(row=row_index, column=index, value=header)
            cell.fill = PatternFill("solid", fgColor="1E3A8A")
            cell.font = Font(color="FFFFFF", bold=True)
            cell.alignment = Alignment(horizontal="center", vertical="center")

        for row in rows:
            row_index += 1
            for column, value in enumerate(row, start=1):
                cell = sheet.cell(row=row_index, column=column, value=value)
                cell.alignment = Alignment(vertical="top", wrap_text=True)

        sheet.freeze_panes = sheet.cell(row=max(len(summary_lines) + 2, 2), column=1)
        sheet.auto_filter.ref = f"A{len(summary_lines) + 1}:{get_column_letter(len(headers))}{max(row_index, len(summary_lines) + 1)}"

        for index, width in enumerate(column_widths, start=1):
            sheet.column_dimensions[get_column_letter(index)].width = width

        buffer = BytesIO()
        workbook.save(buffer)
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    if fmt == "docx":
        from docx import Document
        from docx.enum.section import WD_ORIENTATION
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Mm

        document = Document()
        section = document.sections[0]
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Mm(10)
        section.bottom_margin = Mm(10)
        section.left_margin = Mm(10)
        section.right_margin = Mm(10)
        heading = document.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for line in summary_lines:
            document.add_paragraph(line)

        table = document.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        table.autofit = False

        available = section.page_width - section.left_margin - section.right_margin
        total_width = max(sum(column_widths), 1)
        docx_widths = [int(available * (width / total_width)) for width in column_widths]

        for index, header in enumerate(headers):
            run = table.rows[0].cells[index].paragraphs[0].add_run(header)
            run.bold = True
            table.rows[0].cells[index].width = docx_widths[index]
            table.rows[0].cells[index].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                cells[index].text = value
                cells[index].width = docx_widths[index]

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    font_name = "Helvetica"
    font_path = _resolve_font_path()
    if font_path:
        try:
            pdfmetrics.registerFont(TTFont("CustomFont", font_path))
            font_name = "CustomFont"
        except Exception:
            pass

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("title", parent=styles["Heading1"], fontName=font_name, fontSize=16, alignment=1, spaceAfter=6)
    text_style = ParagraphStyle("text", parent=styles["Normal"], fontName=font_name, fontSize=9, leading=11)
    header_style = ParagraphStyle("header", parent=text_style, alignment=1, textColor=colors.white)
    cell_style = ParagraphStyle("cell", parent=text_style, fontSize=8, leading=10)
    table_data = [[Paragraph(header, header_style) for header in headers]]
    for row in rows:
        table_data.append([Paragraph(value, cell_style) for value in row])

    buffer = BytesIO()
    document = SimpleDocTemplate(
        buffer,
        pagesize=landscape(A4),
        leftMargin=10 * mm,
        rightMargin=10 * mm,
        topMargin=10 * mm,
        bottomMargin=10 * mm,
    )
    elements = [Paragraph(title, title_style)]
    elements.extend(Paragraph(line, text_style) for line in summary_lines)
    elements.append(Spacer(1, 8))
    available_width = landscape(A4)[0] - 20 * mm
    table = Table(table_data, colWidths=_fit_pdf_widths(headers, rows, available_width), repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E3A8A")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("ALIGN", (0, 0), (-1, 0), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#D1D5DB")),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.HexColor("#F3F4F6")]),
            ]
        )
    )
    elements.append(table)
    document.build(elements)
    buffer.seek(0)
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


def _dashboard_section_payload(section: str, dashboard: ReportsDashboard) -> tuple[str, list[str], list[str], list[list[str]], str]:
    generated = _format_ts(dashboard.generated_at)
    if section == "user-actions":
        return (
            "Отчет по действиям пользователей",
            [f"Сформировано: {generated}", f"Активные пользователи: {dashboard.user_actions.active_users}", f"TOTP включен у пользователей: {dashboard.user_actions.totp_enabled_users}"],
            ["Пользователь", "Аудит", "Успешные входы", "Ошибки входа", "Ревью", "Всего"],
            [[item.user_label, str(item.audit_actions), str(item.auth_success), str(item.auth_failures), str(item.review_actions), str(item.total_actions)] for item in dashboard.user_actions.top_users],
            "user-actions-report",
        )
    if section == "groups":
        return (
            "Отчет по группам камер",
            [f"Сформировано: {generated}", f"Групп в отчете: {len(dashboard.groups)}"],
            ["Группа", "Камер", "Онлайн", "Оффлайн", "Событий", "Распознано", "Pending review", "Файлов", "Объем"],
            [[item.name, str(item.camera_count), str(item.online_cameras), str(item.offline_cameras), str(item.event_count), str(item.recognized_count), str(item.pending_reviews), str(item.recordings_count), _format_bytes(item.recordings_size_bytes)] for item in dashboard.groups],
            "group-report",
        )
    if section == "cameras":
        return (
            "Отчет по камерам",
            [f"Сформировано: {generated}", f"Камер в отчете: {len(dashboard.cameras)}"],
            ["Камера", "Группа", "Тип", "Процессор", "Онлайн", "PTZ", "Событий", "Motion", "Unknown", "Архив", "Последнее событие"],
            [[item.name, item.group_name or "-", item.connection_kind, item.assigned_processor or "-", "Да" if item.is_online else "Нет", "Да" if item.supports_ptz else "Нет", str(item.event_count), str(item.motion_count), str(item.unknown_count), _format_bytes(item.recordings_size_bytes), _format_ts(item.last_event_ts)] for item in dashboard.cameras],
            "camera-report",
        )
    if section == "processors":
        return (
            "Отчет по процессорам",
            [f"Сформировано: {generated}", f"Процессоров в отчете: {len(dashboard.processors)}"],
            ["Процессор", "Статус", "IP", "Версия", "Камер", "Событий", "Файлов", "CPU", "RAM", "GPU", "Uptime"],
            [[item.name, "Онлайн" if item.is_online else item.status, item.ip_address or "-", item.version or "-", str(item.assigned_cameras), str(item.event_count), str(item.recordings_count), f"{item.cpu_percent}%" if item.cpu_percent is not None else "-", f"{item.ram_percent}%" if item.ram_percent is not None else "-", f"{item.gpu_util_percent}%" if item.gpu_util_percent is not None else "-", _format_duration(item.uptime_seconds)] for item in dashboard.processors],
            "processor-report",
        )
    if section == "events":
        return (
            "Отчет по событиям и ревью",
            [f"Сформировано: {generated}", f"Всего событий: {dashboard.events.total_events}", f"Pending review: {dashboard.events.pending_reviews}", f"Среднее время ревью: {_format_duration(dashboard.events.average_review_seconds)}"],
            ["Тип события", "Количество"],
            [[item.label, str(item.value)] for item in dashboard.events.events_by_type],
            "events-report",
        )
    if section == "archive":
        return (
            "Отчет по архиву",
            [f"Сформировано: {generated}", f"Всего файлов: {dashboard.archive.total_files}", f"Общий объем: {_format_bytes(dashboard.archive.total_bytes)}"],
            ["Камера", "Файлов", "Объем", "Последняя запись"],
            [[item.camera_name, str(item.file_count), _format_bytes(item.total_bytes), _format_ts(item.last_recording_at)] for item in dashboard.archive.by_camera],
            "archive-report",
        )
    if section == "security":
        return (
            "Отчет по безопасности",
            [f"Сформировано: {generated}", f"Покрытие TOTP: {dashboard.security.totp_coverage_percent:.1f}%", f"Ошибки входа: {dashboard.security.failed_logins}"],
            ["Время", "Пользователь", "Метод", "Причина", "IP"],
            [[_format_ts(item.occurred_at), item.user_label, item.method, item.reason or "-", item.source_ip or "-"] for item in dashboard.security.recent_failures],
            "security-report",
        )
    raise HTTPException(status_code=400, detail="Unknown report section")


@router.get("/dashboard", response_model=ReportsDashboard)
async def reports_dashboard(
    date_from: Optional[str] = Query(default=None),
    date_to: Optional[str] = Query(default=None),
    group_id: Optional[int] = Query(default=None),
    camera_id: Optional[int] = Query(default=None),
    processor_id: Optional[int] = Query(default=None),
    user_id: Optional[int] = Query(default=None),
    session: AsyncSession = Depends(get_session),
    current_user: models.User = Depends(get_current_user),
) -> ReportsDashboard:
    if not is_at_least_user(current_user):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Insufficient permissions")

    date_from_dt = _parse_iso_datetime(date_from)
    date_to_dt = _parse_iso_datetime(date_to)

    users = (await session.execute(select(models.User).order_by(models.User.login.asc()))).scalars().all()
    users_by_id = {user.user_id: user for user in users}
    filtered_users = [user for user in users if user_id is None or user.user_id == user_id]

    groups = (await session.execute(select(models.Group).order_by(models.Group.name.asc()))).scalars().all()
    groups_by_id = {group.group_id: group for group in groups}

    processors = (await session.execute(select(models.Processor).order_by(models.Processor.name.asc()))).scalars().all()
    processors_by_id = {processor.processor_id: processor for processor in processors}

    assignments = (await session.execute(select(models.ProcessorCameraAssignment))).scalars().all()
    camera_to_processor_ids: dict[int, list[int]] = defaultdict(list)
    processor_to_camera_ids: dict[int, list[int]] = defaultdict(list)
    for assignment in assignments:
        camera_to_processor_ids[assignment.camera_id].append(assignment.processor_id)
        processor_to_camera_ids[assignment.processor_id].append(assignment.camera_id)

    cameras = (
        await session.execute(
            select(models.Camera)
            .where(models.Camera.deleted_at.is_(None))
            .order_by(models.Camera.name.asc(), models.Camera.camera_id.asc())
        )
    ).scalars().all()

    visible_cameras: list[models.Camera] = []
    for camera in cameras:
        if group_id is not None and camera.group_id != group_id:
            continue
        if camera_id is not None and camera.camera_id != camera_id:
            continue
        if processor_id is not None and processor_id not in camera_to_processor_ids.get(camera.camera_id, []):
            continue
        visible_cameras.append(camera)

    visible_camera_ids = {camera.camera_id for camera in visible_cameras}
    visible_camera_map = {camera.camera_id: camera for camera in visible_cameras}

    visible_groups = [
        group
        for group in groups
        if ((group_id is None and any(camera.group_id == group.group_id for camera in visible_cameras)) or (group_id is not None and group.group_id == group_id))
    ]

    visible_processors: list[models.Processor] = []
    for processor in processors:
        if processor_id is not None and processor.processor_id != processor_id:
            continue
        if camera_id is not None and camera_id not in processor_to_camera_ids.get(processor.processor_id, []):
            continue
        if group_id is not None:
            assigned_camera_ids = processor_to_camera_ids.get(processor.processor_id, [])
            if not any(visible_camera_map.get(cid) for cid in assigned_camera_ids):
                continue
        visible_processors.append(processor)

    event_type_rows = (await session.execute(select(models.EventType))).scalars().all()
    event_type_map = {event_type.event_type_id: event_type.name for event_type in event_type_rows}

    if visible_camera_ids:
        event_stmt = select(models.Event).where(models.Event.camera_id.in_(visible_camera_ids))
    else:
        event_stmt = select(models.Event).where(models.Event.event_id == -1)
    event_stmt = event_stmt.order_by(models.Event.event_ts.desc())
    if date_from_dt:
        event_stmt = event_stmt.where(models.Event.event_ts >= date_from_dt)
    if date_to_dt:
        event_stmt = event_stmt.where(models.Event.event_ts <= date_to_dt)
    if processor_id is not None:
        event_stmt = event_stmt.where(models.Event.processor_id == processor_id)
    events = (await session.execute(event_stmt)).scalars().all()
    event_ids = [event.event_id for event in events]

    review_rows: list[models.EventReview] = []
    if event_ids:
        review_rows = (
            await session.execute(
                select(models.EventReview)
                .where(models.EventReview.event_id.in_(event_ids))
                .order_by(models.EventReview.updated_at.desc())
            )
        ).scalars().all()
    reviews_by_event_id = {review.event_id: review for review in review_rows}

    auth_stmt = select(models.AuthEvent).order_by(models.AuthEvent.occurred_at.desc())
    if date_from_dt:
        auth_stmt = auth_stmt.where(models.AuthEvent.occurred_at >= date_from_dt)
    if date_to_dt:
        auth_stmt = auth_stmt.where(models.AuthEvent.occurred_at <= date_to_dt)
    if user_id is not None:
        auth_stmt = auth_stmt.where(models.AuthEvent.user_id == user_id)
    auth_events = (await session.execute(auth_stmt)).scalars().all()

    audit_stmt = select(models.AuditLog).order_by(models.AuditLog.changed_at.desc())
    if date_from_dt:
        audit_stmt = audit_stmt.where(models.AuditLog.changed_at >= date_from_dt)
    if date_to_dt:
        audit_stmt = audit_stmt.where(models.AuditLog.changed_at <= date_to_dt)
    if user_id is not None:
        audit_stmt = audit_stmt.where(models.AuditLog.changed_by == user_id)
    audit_logs = (await session.execute(audit_stmt)).scalars().all()

    recording_rows = (
        await session.execute(
            select(models.RecordingFile, models.VideoStream, models.StorageTarget)
            .join(models.VideoStream, models.RecordingFile.video_stream_id == models.VideoStream.video_stream_id)
            .join(models.StorageTarget, models.RecordingFile.storage_target_id == models.StorageTarget.storage_target_id)
            .order_by(models.RecordingFile.started_at.desc())
        )
    ).all()

    visible_recordings = []
    for recording, stream, storage_target in recording_rows:
        if stream.camera_id not in visible_camera_ids:
            continue
        if date_from_dt and recording.started_at < date_from_dt:
            continue
        if date_to_dt and recording.started_at > date_to_dt:
            continue
        visible_recordings.append((recording, stream, storage_target))

    mfa_rows = (
        await session.execute(
            select(models.UserMfaMethod).where(
                models.UserMfaMethod.mfa_type == "totp",
                models.UserMfaMethod.is_enabled.is_(True),
            )
        )
    ).scalars().all()
    totp_user_ids = {row.user_id for row in mfa_rows}
    api_keys = (await session.execute(select(models.ApiKey))).scalars().all()

    user_stat_map: dict[Optional[int], dict[str, int]] = defaultdict(lambda: defaultdict(int))
    for item in audit_logs:
        user_stat_map[item.changed_by]["audit_actions"] += 1
        user_stat_map[item.changed_by]["total_actions"] += 1
    for item in auth_events:
        if item.success:
            user_stat_map[item.user_id]["auth_success"] += 1
        else:
            user_stat_map[item.user_id]["auth_failures"] += 1
        user_stat_map[item.user_id]["total_actions"] += 1
    for review in review_rows:
        if review.reviewer_user_id is not None:
            user_stat_map[review.reviewer_user_id]["review_actions"] += 1
            user_stat_map[review.reviewer_user_id]["total_actions"] += 1

    top_users = sorted(
        [
            UserActionActorStat(
                user_id=user_key,
                user_label=_user_label(users_by_id.get(user_key), user_key),
                audit_actions=stats.get("audit_actions", 0),
                auth_success=stats.get("auth_success", 0),
                auth_failures=stats.get("auth_failures", 0),
                review_actions=stats.get("review_actions", 0),
                total_actions=stats.get("total_actions", 0),
            )
            for user_key, stats in user_stat_map.items()
            if user_id is None or user_key == user_id
        ],
        key=lambda item: (-item.total_actions, item.user_label),
    )[:10]

    recent_actions: list[RecentUserAction] = []
    for item in audit_logs[:10]:
        recent_actions.append(
            RecentUserAction(
                action_kind="audit",
                occurred_at=item.changed_at.isoformat(),
                user_id=item.changed_by,
                user_label=_user_label(users_by_id.get(item.changed_by), item.changed_by),
                action=f"{item.action} {item.table_name}",
                details=item.change_data,
                source_ip=item.source_ip,
            )
        )
    for item in auth_events[:10]:
        recent_actions.append(
            RecentUserAction(
                action_kind="auth",
                occurred_at=item.occurred_at.isoformat(),
                user_id=item.user_id,
                user_label=_user_label(users_by_id.get(item.user_id), item.user_id),
                action=f"Авторизация ({item.method})",
                details=item.reason,
                success=item.success,
                source_ip=item.source_ip,
            )
        )
    for item in review_rows[:10]:
        recent_actions.append(
            RecentUserAction(
                action_kind="review",
                occurred_at=item.updated_at.isoformat(),
                user_id=item.reviewer_user_id,
                user_label=_user_label(users_by_id.get(item.reviewer_user_id), item.reviewer_user_id),
                action=f"Ревью: {item.status}",
                details=item.note,
            )
        )
    recent_actions.sort(key=lambda item: item.occurred_at, reverse=True)
    recent_actions = recent_actions[:15]

    events_by_camera: dict[int, list[models.Event]] = defaultdict(list)
    events_by_processor: dict[Optional[int], list[models.Event]] = defaultdict(list)
    event_type_counter: Counter[str] = Counter()
    recognized_events = 0
    unknown_events = 0
    motion_events = 0
    person_events = 0
    for event in events:
        events_by_camera[event.camera_id].append(event)
        events_by_processor[event.processor_id].append(event)
        event_name = event_type_map.get(event.event_type_id, f"type:{event.event_type_id}")
        event_type_counter[event_name] += 1
        if event_name == "face_recognized":
            recognized_events += 1
        elif event_name == "face_unknown":
            unknown_events += 1
        elif "motion" in event_name:
            motion_events += 1
        elif "person" in event_name:
            person_events += 1

    pending_reviews = sum(1 for review in review_rows if review.status == "pending")
    approved_reviews = sum(1 for review in review_rows if review.status == "approved")
    rejected_reviews = sum(1 for review in review_rows if review.status == "rejected")

    reviewer_counter: dict[Optional[int], Counter[str]] = defaultdict(Counter)
    review_durations: list[float] = []
    for review in review_rows:
        reviewer_counter[review.reviewer_user_id][review.status] += 1
        if review.status != "pending":
            review_durations.append(max((review.updated_at - review.created_at).total_seconds(), 0))

    recordings_by_camera: dict[int, dict[str, object]] = defaultdict(lambda: {"count": 0, "bytes": 0, "last": None})
    recordings_by_storage: dict[int, dict[str, object]] = defaultdict(lambda: {"count": 0, "bytes": 0, "name": ""})
    video_files = 0
    snapshot_files = 0
    for recording, stream, storage_target in visible_recordings:
        camera_bucket = recordings_by_camera[stream.camera_id]
        camera_bucket["count"] = int(camera_bucket["count"]) + 1
        camera_bucket["bytes"] = int(camera_bucket["bytes"]) + _safe_int(recording.file_size_bytes)
        last_seen = camera_bucket["last"]
        if last_seen is None or (recording.started_at and recording.started_at > last_seen):
            camera_bucket["last"] = recording.started_at

        storage_bucket = recordings_by_storage[storage_target.storage_target_id]
        storage_bucket["count"] = int(storage_bucket["count"]) + 1
        storage_bucket["bytes"] = int(storage_bucket["bytes"]) + _safe_int(recording.file_size_bytes)
        storage_bucket["name"] = storage_target.name

        if recording.file_kind == "video":
            video_files += 1
        elif recording.file_kind == "snapshot":
            snapshot_files += 1

    user_actions_report = UserActionsReport(
        active_users=len(
            {item.user_id for item in auth_events if item.user_id is not None}
            | {item.changed_by for item in audit_logs if item.changed_by is not None}
            | {item.reviewer_user_id for item in review_rows if item.reviewer_user_id is not None}
        ),
        total_audit_actions=len(audit_logs),
        total_auth_events=len(auth_events),
        failed_auth_events=sum(1 for item in auth_events if not item.success),
        review_actions=sum(1 for item in review_rows if item.reviewer_user_id is not None),
        totp_enabled_users=sum(1 for user in filtered_users if user.user_id in totp_user_ids),
        top_users=top_users,
        recent_actions=recent_actions,
    )

    group_items: list[GroupReportItem] = []
    for group in visible_groups:
        group_camera_ids = [camera.camera_id for camera in visible_cameras if camera.group_id == group.group_id]
        online_camera_ids = sum(
            1
            for cid in group_camera_ids
            if any(_is_processor_online(processors_by_id[pid]) for pid in camera_to_processor_ids.get(cid, []) if pid in processors_by_id)
        )
        group_items.append(
            GroupReportItem(
                group_id=group.group_id,
                name=group.name,
                camera_count=len(group_camera_ids),
                online_cameras=online_camera_ids,
                offline_cameras=max(len(group_camera_ids) - online_camera_ids, 0),
                event_count=sum(len(events_by_camera.get(cid, [])) for cid in group_camera_ids),
                recognized_count=sum(
                    1
                    for cid in group_camera_ids
                    for event in events_by_camera.get(cid, [])
                    if event_type_map.get(event.event_type_id) == "face_recognized"
                ),
                pending_reviews=sum(
                    1
                    for cid in group_camera_ids
                    for event in events_by_camera.get(cid, [])
                    if reviews_by_event_id.get(event.event_id) and reviews_by_event_id[event.event_id].status == "pending"
                ),
                recordings_count=sum(int(recordings_by_camera[cid]["count"]) for cid in group_camera_ids),
                recordings_size_bytes=sum(int(recordings_by_camera[cid]["bytes"]) for cid in group_camera_ids),
            )
        )

    camera_items: list[CameraReportItem] = []
    for camera in visible_cameras:
        processor_names = [
            processors_by_id[pid].name
            for pid in camera_to_processor_ids.get(camera.camera_id, [])
            if pid in processors_by_id
        ]
        camera_events = events_by_camera.get(camera.camera_id, [])
        last_event = max((event.event_ts for event in camera_events), default=None)
        camera_items.append(
            CameraReportItem(
                camera_id=camera.camera_id,
                name=camera.name,
                location=camera.location,
                group_name=groups_by_id[camera.group_id].name if camera.group_id in groups_by_id else None,
                connection_kind=camera.connection_kind,
                assigned_processor=", ".join(sorted(dict.fromkeys(processor_names))) if processor_names else None,
                detection_enabled=camera.detection_enabled,
                supports_ptz=camera.supports_ptz,
                is_online=any(
                    _is_processor_online(processors_by_id[pid])
                    for pid in camera_to_processor_ids.get(camera.camera_id, [])
                    if pid in processors_by_id
                ),
                event_count=len(camera_events),
                recognized_count=sum(1 for event in camera_events if event_type_map.get(event.event_type_id) == "face_recognized"),
                unknown_count=sum(1 for event in camera_events if event_type_map.get(event.event_type_id) == "face_unknown"),
                motion_count=sum(1 for event in camera_events if "motion" in event_type_map.get(event.event_type_id, "")),
                pending_reviews=sum(
                    1
                    for event in camera_events
                    if reviews_by_event_id.get(event.event_id) and reviews_by_event_id[event.event_id].status == "pending"
                ),
                recordings_count=int(recordings_by_camera[camera.camera_id]["count"]),
                recordings_size_bytes=int(recordings_by_camera[camera.camera_id]["bytes"]),
                last_event_ts=_format_iso(last_event),
            )
        )

    processor_items: list[ProcessorReportItem] = []
    for processor in visible_processors:
        metrics = _load_metrics(processor)
        assigned_camera_ids = [cid for cid in processor_to_camera_ids.get(processor.processor_id, []) if cid in visible_camera_ids]
        processor_items.append(
            ProcessorReportItem(
                processor_id=processor.processor_id,
                name=processor.name,
                status=processor.status,
                is_online=_is_processor_online(processor),
                ip_address=processor.ip_address,
                version=processor.version,
                last_heartbeat=_format_iso(processor.last_heartbeat),
                assigned_cameras=len(assigned_camera_ids),
                event_count=len(events_by_processor.get(processor.processor_id, [])),
                recordings_count=sum(int(recordings_by_camera[cid]["count"]) for cid in assigned_camera_ids),
                cpu_percent=metrics.get("cpu_percent"),
                ram_percent=metrics.get("ram_percent"),
                gpu_util_percent=metrics.get("gpu_util_percent"),
                uptime_seconds=metrics.get("uptime_seconds"),
            )
        )

    top_reviewers = sorted(
        [
            ReviewerStat(
                user_id=reviewer_id,
                user_label=_user_label(users_by_id.get(reviewer_id), reviewer_id),
                approved=stats.get("approved", 0),
                rejected=stats.get("rejected", 0),
                pending=stats.get("pending", 0),
                total=sum(stats.values()),
            )
            for reviewer_id, stats in reviewer_counter.items()
            if user_id is None or reviewer_id == user_id
        ],
        key=lambda item: (-item.total, item.user_label),
    )[:10]

    archive_by_camera = sorted(
        [
            ArchiveCameraStat(
                camera_id=camera.camera_id,
                camera_name=camera.name,
                file_count=int(recordings_by_camera[camera.camera_id]["count"]),
                total_bytes=int(recordings_by_camera[camera.camera_id]["bytes"]),
                last_recording_at=_format_iso(recordings_by_camera[camera.camera_id]["last"]),
            )
            for camera in visible_cameras
            if int(recordings_by_camera[camera.camera_id]["count"]) > 0
        ],
        key=lambda item: (-item.total_bytes, item.camera_name),
    )

    archive_by_storage = sorted(
        [
            ReportStorageStat(
                storage_target_id=storage_id,
                name=str(bucket["name"] or f"Хранилище {storage_id}"),
                file_count=int(bucket["count"]),
                total_bytes=int(bucket["bytes"]),
            )
            for storage_id, bucket in recordings_by_storage.items()
        ],
        key=lambda item: (-item.total_bytes, item.name),
    )

    security_failures = [
        SecurityFailureItem(
            occurred_at=item.occurred_at.isoformat(),
            user_id=item.user_id,
            user_label=_user_label(users_by_id.get(item.user_id), item.user_id),
            method=item.method,
            reason=item.reason,
            source_ip=item.source_ip,
        )
        for item in auth_events
        if not item.success
    ][:10]

    successful_logins = sum(1 for item in auth_events if item.success)
    failed_logins = sum(1 for item in auth_events if not item.success)
    totp_enabled_filtered = sum(1 for user in filtered_users if user.user_id in totp_user_ids)

    return ReportsDashboard(
        generated_at=datetime.utcnow().isoformat(),
        date_from=date_from,
        date_to=date_to,
        group_id=group_id,
        camera_id=camera_id,
        processor_id=processor_id,
        user_id=user_id,
        user_actions=user_actions_report,
        groups=sorted(group_items, key=lambda item: (-item.event_count, item.name)),
        cameras=sorted(camera_items, key=lambda item: (-item.event_count, item.name)),
        processors=sorted(processor_items, key=lambda item: (-item.assigned_cameras, item.name)),
        events=EventReviewReport(
            total_events=len(events),
            recognized_events=recognized_events,
            unknown_events=unknown_events,
            motion_events=motion_events,
            person_events=person_events,
            pending_reviews=pending_reviews,
            approved_reviews=approved_reviews,
            rejected_reviews=rejected_reviews,
            average_review_seconds=round(sum(review_durations) / len(review_durations), 2) if review_durations else None,
            events_by_type=[ReportValueLabel(label=label, value=value) for label, value in event_type_counter.most_common()],
            top_reviewers=top_reviewers,
        ),
        archive=ArchiveReport(
            total_files=len(visible_recordings),
            total_bytes=sum(_safe_int(recording.file_size_bytes) for recording, _, _ in visible_recordings),
            video_files=video_files,
            snapshot_files=snapshot_files,
            by_camera=archive_by_camera,
            by_storage=archive_by_storage,
        ),
        security=SecurityReport(
            total_users=len(filtered_users),
            totp_enabled_users=totp_enabled_filtered,
            totp_coverage_percent=round((totp_enabled_filtered / len(filtered_users)) * 100, 2) if filtered_users else 0,
            api_keys_total=len(api_keys),
            api_keys_active=sum(1 for key in api_keys if key.is_active),
            successful_logins=successful_logins,
            failed_logins=failed_logins,
            recent_failures=security_failures,
        ),
    )


@router.get("/export")
async def export_dashboard_section(
    section: str = Query(..., description="user-actions|groups|cameras|processors|events|archive|security"),
    format: str = Query(default="pdf", description="pdf|xlsx|docx"),
    date_from: Optional[str] = Query(default=None),
    date_to: Optional[str] = Query(default=None),
    group_id: Optional[int] = Query(default=None),
    camera_id: Optional[int] = Query(default=None),
    processor_id: Optional[int] = Query(default=None),
    user_id: Optional[int] = Query(default=None),
    session: AsyncSession = Depends(get_session),
    current_user: models.User = Depends(get_current_user),
):
    fmt = format.lower().strip()
    if fmt not in {"pdf", "xlsx", "docx"}:
        raise HTTPException(status_code=400, detail="format must be pdf, xlsx or docx")

    dashboard = await reports_dashboard(
        date_from=date_from,
        date_to=date_to,
        group_id=group_id,
        camera_id=camera_id,
        processor_id=processor_id,
        user_id=user_id,
        session=session,
        current_user=current_user,
    )
    title, summary_lines, headers, rows, prefix = _dashboard_section_payload(section, dashboard)
    return _render_export_table(
        title=title,
        summary_lines=summary_lines,
        headers=headers,
        rows=rows,
        fmt=fmt,
        filename_prefix=prefix,
    )


@router.get("/appearances", response_model=AppearanceReport)
async def appearances_report(
    date_from: Optional[str] = Query(default=None, description="ISO datetime start"),
    date_to: Optional[str] = Query(default=None, description="ISO datetime end"),
    person_id: Optional[int] = Query(default=None),
    session: AsyncSession = Depends(get_session),
    current_user: models.User = Depends(get_current_user),
) -> AppearanceReport:
    if not is_at_least_user(current_user):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Insufficient permissions")

    items, _, _, _ = await _load_appearance_items(
        date_from=date_from,
        date_to=date_to,
        person_id=person_id,
        session=session,
    )

    return AppearanceReport(
        date_from=date_from,
        date_to=date_to,
        person_id=person_id,
        total=len(items),
        items=items,
    )


@router.get("/appearances/export")
async def appearances_report_export(
    format: str = Query(default="pdf", description="pdf|xlsx|docx"),
    date_from: Optional[str] = Query(default=None, description="ISO datetime start"),
    date_to: Optional[str] = Query(default=None, description="ISO datetime end"),
    person_id: Optional[int] = Query(default=None),
    session: AsyncSession = Depends(get_session),
    current_user: models.User = Depends(get_current_user),
):
    if not is_at_least_user(current_user):
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Insufficient permissions")

    fmt = format.lower().strip()
    if fmt not in {"pdf", "xlsx", "docx"}:
        raise HTTPException(status_code=400, detail="format must be pdf, xlsx or docx")

    items, person_label, date_from_value, date_to_value = await _load_appearance_items(
        date_from=date_from,
        date_to=date_to,
        person_id=person_id,
        session=session,
    )

    period = _format_period(date_from_value, date_to_value)
    subject = person_label or "Все персоны"
    generated_at = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
    headers = ["№", "Время", "Камера", "Локация", "Группа", "Персона", "Уверенность"]
    return _render_export_table(
        title="Отчет по появлениям",
        summary_lines=[
            f"Период: {period}",
            f"Персона: {subject}",
            f"Сформировано: {generated_at}",
        ],
        headers=headers,
        rows=[_appearance_row(index, item) for index, item in enumerate(items, start=1)],
        fmt=fmt,
        filename_prefix="appearance-report",
    )
